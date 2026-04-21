"""Cover letter generation API routes."""
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from backend.database import get_db
from backend.schemas.resume import CoverLetterRequest, CoverLetterResponse
from backend.utils.auth import get_current_user
from backend.services.cover_letter_service import cover_letter_service
from backend.services.memory_service import memory_service

router = APIRouter(prefix="/api/cover-letter", tags=["Cover Letter"])


from fastapi.responses import FileResponse
from backend.schemas.resume import CoverLetterRequest, CoverLetterResponse, CoverLetterDownloadRequest
from backend.services.file_service import file_service
from datetime import datetime
import os
import uuid
from backend.config import settings

@router.post("/generate", response_model=CoverLetterResponse)
def generate_cover_letter(
    req: CoverLetterRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    try:
        profile = memory_service.get_user_context(db, current_user["id"])
        result = cover_letter_service.generate(
            company=req.company_name,
            role=req.role,
            jd=req.job_description or "",
            skills=req.key_skills,
            tone=req.tone,
            context=req.additional_context or "",
            profile=profile,
        )
        return CoverLetterResponse(**result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@router.post("/download/{format}")
def download_cover_letter(
    format: str,
    req: CoverLetterDownloadRequest,
    current_user: dict = Depends(get_current_user),
):
    content = req.content
    content["current_date"] = datetime.today().strftime('%B %d, %Y')
    
    if format == "pdf":
        html = file_service.render_template("cover_letter_reference.html", content)
        filepath = file_service.generate_pdf(html)
        if filepath and filepath.endswith(".html"):
            return FileResponse(filepath, filename="cover_letter.html", media_type="text/html")
    elif format == "html":
        html = file_service.render_template("cover_letter_reference.html", content)
        filepath = os.path.join(settings.GENERATED_DIR, f"cover_letter_{uuid.uuid4().hex[:8]}.html")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html)
    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Contact block
            name = content.get("user_name", "Your Name")
            contact = [content.get("user_location", ""), content.get("user_phone", ""), content.get("user_email", "")]
            doc.add_paragraph(name).bold = True
            doc.add_paragraph(" | ".join([c for c in contact if c]))
            doc.add_paragraph("")
            
            doc.add_paragraph(content["current_date"])
            doc.add_paragraph("")
            
            if content.get("recipient_name"): doc.add_paragraph(content["recipient_name"] + (f", {content['recipient_title']}" if content.get("recipient_title") else ""))
            if content.get("company_name"): doc.add_paragraph(content["company_name"])
            if content.get("company_address"): doc.add_paragraph(content["company_address"])
            
            doc.add_paragraph("")
            doc.add_paragraph(content.get("salutation", "Dear Hiring Manager,"))
            doc.add_paragraph("")
            
            for p in content.get("body_paragraphs", []):
                doc.add_paragraph(p)
                doc.add_paragraph("")
                
            doc.add_paragraph(content.get("sign_off", "Best regards,"))
            doc.add_paragraph(name)
            
            filepath = os.path.join(settings.GENERATED_DIR, f"cover_letter_{uuid.uuid4().hex[:8]}.docx")
            doc.save(filepath)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX failure: {str(e)}")
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use pdf, html, or docx")
        
    if not filepath:
        raise HTTPException(status_code=500, detail="File generation failed")
        
    if filepath.endswith(".pdf"):
        media_type = "application/pdf"
        dl_name = "cover_letter.pdf"
    elif filepath.endswith(".html"):
        media_type = "text/html"
        dl_name = "cover_letter.html"
    elif filepath.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        dl_name = "cover_letter.docx"
    else:
        media_type = "application/octet-stream"
        dl_name = f"cover_letter.{format}"

    return FileResponse(filepath, filename=dl_name, media_type=media_type)
