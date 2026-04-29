"""File generation service — HTML-to-PDF and DOCX generation."""
import os
import uuid
import logging
from pathlib import Path
from typing import Optional
from bs4 import BeautifulSoup
from jinja2 import Environment, FileSystemLoader
from backend.config import settings

logger = logging.getLogger(__name__)

jinja_env = Environment(
    loader=FileSystemLoader(settings.TEMPLATE_DIR),
    autoescape=False,  # Disable autoescape for resume HTML rendering
)

# Max file sizes (in bytes)
MAX_PDF_SIZE = 5 * 1024 * 1024  # 5MB
MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10MB
MAX_TXT_SIZE = 2 * 1024 * 1024  # 2MB


class FileService:
    @staticmethod
    def _pdf_escape(text: str) -> str:
        return (
            str(text or "")
            .replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )

    def _generate_simple_pdf(self, text_content: str, filepath: str) -> str:
        """Generate a basic text PDF without external system libraries."""
        lines = []
        for raw_line in str(text_content or "").splitlines():
            stripped = raw_line.strip()
            if not stripped:
                lines.append("")
                continue

            words = stripped.split()
            current = []
            current_len = 0
            for word in words:
                projected = current_len + len(word) + (1 if current else 0)
                if projected > 95:
                    lines.append(" ".join(current))
                    current = [word]
                    current_len = len(word)
                else:
                    current.append(word)
                    current_len = projected
            if current:
                lines.append(" ".join(current))

        if not lines:
            lines = ["Resume"]

        pages = []
        line_height = 14
        max_lines_per_page = 48
        for start in range(0, len(lines), max_lines_per_page):
            pages.append(lines[start:start + max_lines_per_page])

        objects = []

        def add_object(payload: str) -> int:
            objects.append(payload)
            return len(objects)

        font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        page_ids = []

        for page_lines in pages:
            content_commands = ["BT", "/F1 11 Tf", "72 770 Td"]
            first_line = True
            for line in page_lines:
                escaped = self._pdf_escape(line)
                if first_line:
                    content_commands.append(f"({escaped}) Tj")
                    first_line = False
                else:
                    content_commands.append(f"0 -{line_height} Td")
                    content_commands.append(f"({escaped}) Tj")
            content_commands.append("ET")
            content_stream = "\n".join(content_commands)
            content_id = add_object(
                f"<< /Length {len(content_stream.encode('utf-8'))} >>\nstream\n{content_stream}\nendstream"
            )
            page_id = add_object(
                f"<< /Type /Page /Parent {{PAGES_ID}} 0 R /MediaBox [0 0 595 842] "
                f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>"
            )
            page_ids.append(page_id)

        kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
        pages_id = add_object(f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>")

        for idx, payload in enumerate(objects):
            if "{PAGES_ID}" in payload:
                objects[idx] = payload.replace("{PAGES_ID}", str(pages_id))

        catalog_id = add_object(f"<< /Type /Catalog /Pages {pages_id} 0 R >>")

        pdf_parts = ["%PDF-1.4\n"]
        offsets = [0]
        current_offset = len(pdf_parts[0].encode("utf-8"))

        for object_id, payload in enumerate(objects, start=1):
            obj_str = f"{object_id} 0 obj\n{payload}\nendobj\n"
            offsets.append(current_offset)
            pdf_parts.append(obj_str)
            current_offset += len(obj_str.encode("utf-8"))

        xref_offset = current_offset
        xref_lines = [f"xref\n0 {len(objects) + 1}\n", "0000000000 65535 f \n"]
        for offset in offsets[1:]:
            xref_lines.append(f"{offset:010d} 00000 n \n")
        trailer = (
            f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF"
        )
        pdf_parts.extend(xref_lines)
        pdf_parts.append(trailer)

        with open(filepath, "wb") as f:
            f.write("".join(pdf_parts).encode("utf-8"))

        return filepath

    def _generate_pdf_fallback_from_html(self, html_content: str, filepath: str) -> str:
        text_content = BeautifulSoup(html_content, "html.parser").get_text("\n")
        return self._generate_simple_pdf(text_content, filepath)

    def render_template(self, template_name: str, context: dict) -> str:
        """Render an HTML template with context data."""
        try:
            # Normalize skills format: support both dict-of-lists and dict-of-strings
            if "skills" in context and isinstance(context["skills"], dict):
                normalized = {}
                for cat, items in context["skills"].items():
                    if isinstance(items, list):
                        normalized[cat] = ", ".join(items)
                    else:
                        normalized[cat] = str(items)
                context["skills"] = normalized

            template = jinja_env.get_template(template_name)
            return template.render(**context)
        except Exception as e:
            logger.error(f"Template rendering error: {e}")
            return f"<html><body><p>Template error: {e}</p></body></html>"

    def _validate_file_size(self, filepath: str, max_size: int, file_type: str) -> bool:
        """Check if file exceeds max size and log appropriately."""
        try:
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return False
            
            file_size = os.path.getsize(filepath)
            if file_size > max_size:
                logger.warning(
                    f"{file_type} generation warning: File size {file_size} bytes exceeds limit {max_size} bytes",
                    extra={
                        "file_path": filepath,
                        "file_size": file_size,
                        "max_size": max_size,
                        "file_type": file_type
                    }
                )
                return False
            
            logger.debug(f"{file_type} file validated: {file_size} bytes")
            return True
        except Exception as e:
            logger.error(f"Size validation error: {e}")
            return False

    def generate_pdf(self, html_content: str, filename: Optional[str] = None) -> str:
        """Generate PDF from HTML content. Returns file path."""
        if not filename:
            filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        try:
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(filepath)
            
            # Validate file size
            if not self._validate_file_size(filepath, MAX_PDF_SIZE, "PDF"):
                if os.path.exists(filepath):
                    os.remove(filepath)
                raise ValueError(f"Generated PDF exceeds {MAX_PDF_SIZE} bytes limit")
            
            logger.info(f"PDF generated successfully: {filepath}")
            return filepath
        except ImportError:
            logger.warning("weasyprint not installed. Trying pdfkit fallback...")
            # Try pdfkit as second fallback
            try:
                import pdfkit
                pdfkit.from_string(html_content, filepath)
                
                # Validate file size
                if not self._validate_file_size(filepath, MAX_PDF_SIZE, "PDF"):
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    raise ValueError(f"Generated PDF exceeds {MAX_PDF_SIZE} bytes limit or creation failed")
                
                logger.info(f"PDF generated via pdfkit: {filepath}")
                return filepath
            except (ImportError, Exception) as e2:
                logger.warning(f"pdfkit also failed: {e2}. Using built-in text PDF fallback.")
                fallback_path = self._generate_pdf_fallback_from_html(html_content, filepath)
                logger.info(f"Fallback PDF generated: {fallback_path}")
                return fallback_path
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            fallback_path = self._generate_pdf_fallback_from_html(html_content, filepath)
            logger.info(f"Fallback PDF generated due to error: {fallback_path}")
            return fallback_path

    def generate_docx(self, resume_data: dict, filename: Optional[str] = None) -> str:
        """Generate DOCX from resume data. Returns file path."""
        if not filename:
            filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style setup
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Name
            name = resume_data.get("full_name", "Your Name")
            heading = doc.add_heading(name, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact
            contact = resume_data.get("contact", {})
            contact_parts = []
            for key in ["email", "phone", "linkedin", "github", "portfolio", "location", "leetcode"]:
                val = contact.get(key)
                if val:
                    contact_parts.append(val)
            if contact_parts:
                p = doc.add_paragraph(" | ".join(contact_parts))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            if resume_data.get("summary"):
                doc.add_heading("Professional Summary", level=1)
                summary = resume_data["summary"]
                if isinstance(summary, list):
                    for s in summary:
                        doc.add_paragraph(s)
                else:
                    doc.add_paragraph(str(summary))

            # Education
            education = resume_data.get("education", [])
            if education:
                doc.add_heading("Education", level=1)
                for edu in education:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{edu.get('school', '')} — {edu.get('degree', '')}")
                    run.bold = True
                    details = []
                    if edu.get("dates"):
                        details.append(edu["dates"])
                    if edu.get("grade"):
                        details.append(edu["grade"])
                    if details:
                        p.add_run(f"\n{' | '.join(details)}")

            # Technical Skills
            skills = resume_data.get("skills", {})
            if skills:
                doc.add_heading("Technical Skills", level=1)
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if skill_list:
                            if isinstance(skill_list, list):
                                skill_str = ", ".join(skill_list)
                            else:
                                skill_str = str(skill_list)
                            p = doc.add_paragraph()
                            run = p.add_run(f"{category}: ")
                            run.bold = True
                            p.add_run(skill_str)
                elif isinstance(skills, list):
                    doc.add_paragraph(", ".join(skills))
                else:
                    doc.add_paragraph(str(skills))

            # Experience
            experience = resume_data.get("experience", [])
            if experience:
                doc.add_heading("Experience", level=1)
                for exp in experience:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{exp.get('company', '')} — {exp.get('title', '')}")
                    run.bold = True
                    sub_parts = []
                    if exp.get("location"):
                        sub_parts.append(exp["location"])
                    if exp.get("dates"):
                        sub_parts.append(exp["dates"])
                    if sub_parts:
                        p.add_run(f"\n{' | '.join(sub_parts)}")
                    for bullet in exp.get("bullets", []):
                        doc.add_paragraph(bullet, style='List Bullet')

            # Projects
            projects = resume_data.get("projects", [])
            if projects:
                doc.add_heading("Projects", level=1)
                for proj in projects:
                    p = doc.add_paragraph()
                    run = p.add_run(proj.get("name", ""))
                    run.bold = True
                    tech = proj.get("tech_stack", "")
                    if tech:
                        p.add_run(f" | {tech}")
                    
                    link_line = []
                    if proj.get("live_url"):
                        link_line.append(f"Live: {proj['live_url']}")
                    if proj.get("repo_url"):
                        link_line.append(f"Repo: {proj['repo_url']}")
                    if link_line:
                        doc.add_paragraph(" | ".join(link_line))
                        
                    for bullet in proj.get("bullets", []):
                        doc.add_paragraph(bullet, style='List Bullet')

            # Certifications
            certs = resume_data.get("certifications", [])
            if certs:
                doc.add_heading("Certifications", level=1)
                for cert in certs:
                    parts = [cert.get("name", "")]
                    if cert.get("issuer"):
                        parts.append(cert["issuer"])
                    if cert.get("date"):
                        parts.append(cert["date"])
                    doc.add_paragraph(" — ".join(parts), style='List Bullet')

            # Achievements
            achievements = resume_data.get("achievements", [])
            if achievements:
                doc.add_heading("Achievements", level=1)
                for ach in achievements:
                    doc.add_paragraph(ach, style='List Bullet')

            doc.save(filepath)
            
            # Validate file size
            if not self._validate_file_size(filepath, MAX_DOCX_SIZE, "DOCX"):
                os.remove(filepath)
                raise ValueError(f"Generated DOCX exceeds {MAX_DOCX_SIZE} bytes limit")
            
            logger.info(f"DOCX generated successfully: {filepath}")
            return filepath

        except ImportError:
            logger.error("python-docx not installed")
            return ""
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return ""

    def generate_txt(self, resume_data: dict, filename: Optional[str] = None) -> str:
        """Generate plain text resume."""
        if not filename:
            filename = f"resume_{uuid.uuid4().hex[:8]}.txt"
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        try:
            lines = []
            lines.append(resume_data.get("full_name", "Your Name").upper())
            contact = resume_data.get("contact", {})
            contact_parts = [v for v in contact.values() if v]
            lines.append(" | ".join(contact_parts))
            lines.append("=" * 60)

            if resume_data.get("summary"):
                lines.append("\nPROFESSIONAL SUMMARY")
                lines.append("-" * 40)
                summary = resume_data["summary"]
                if isinstance(summary, list):
                    for s in summary:
                        lines.append(s)
                else:
                    lines.append(str(summary))

            if resume_data.get("education"):
                lines.append("\nEDUCATION")
                lines.append("-" * 40)
                for edu in resume_data.get("education", []):
                    lines.append(f"{edu.get('degree', '')} — {edu.get('school', '')} ({edu.get('dates', '')})")
                    if edu.get("grade"):
                        lines.append(f"  {edu['grade']}")

            skills = resume_data.get("skills", {})
            if skills:
                lines.append("\nTECHNICAL SKILLS")
                lines.append("-" * 40)
                if isinstance(skills, dict):
                    for cat, slist in skills.items():
                        if slist:
                            if isinstance(slist, list):
                                lines.append(f"{cat}: {', '.join(slist)}")
                            else:
                                lines.append(f"{cat}: {slist}")

            for exp in resume_data.get("experience", []):
                lines.append(f"\n{exp.get('company', '')} — {exp.get('title', '')}")
                lines.append(f"{exp.get('location', '')} | {exp.get('dates', '')}")
                for bullet in exp.get("bullets", []):
                    lines.append(f"  • {bullet}")

            if resume_data.get("projects"):
                lines.append("\nPROJECTS")
                lines.append("-" * 40)
                for proj in resume_data.get("projects", []):
                    tech = f" ({proj.get('tech_stack', '')})" if proj.get('tech_stack') else ""
                    lines.append(f"{proj.get('name', '')}{tech}")
                    if proj.get("live_url"):
                        lines.append(f"  Live: {proj['live_url']}")
                    for bullet in proj.get("bullets", []):
                        lines.append(f"  • {bullet}")

            if resume_data.get("certifications"):
                lines.append("\nCERTIFICATIONS")
                lines.append("-" * 40)
                for cert in resume_data.get("certifications", []):
                    parts = [cert.get("name", "")]
                    if cert.get("issuer"):
                        parts.append(cert["issuer"])
                    lines.append(" — ".join(parts))

            if resume_data.get("achievements"):
                lines.append("\nACHIEVEMENTS")
                lines.append("-" * 40)
                for ach in resume_data.get("achievements", []):
                    lines.append(f"  • {ach}")

            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            
            # Validate file size
            if not self._validate_file_size(filepath, MAX_TXT_SIZE, "TXT"):
                os.remove(filepath)
                raise ValueError(f"Generated TXT exceeds {MAX_TXT_SIZE} bytes limit")
            
            logger.info(f"TXT generated successfully: {filepath}")
            return filepath
        
        except Exception as e:
            logger.error(f"TXT generation error: {e}")
            return ""


file_service = FileService()
